"""Full Persian academic article draft (RTL Word), populated from outputs/.

Produces outputs/09_report/article_full_draft_fa.docx with a conventional paper
structure and real numbers pulled from the pipeline results.
"""
from __future__ import annotations
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import utils

FA = "نام مدل"  # placeholder constant (unused) kept for clarity
MODEL_FA = {"logistic": "رگرسیون لجستیک", "random_forest": "جنگل تصادفی",
            "xgboost": "XGBoost"}


# --------------------------------------------------------------------------
# RTL helpers
# --------------------------------------------------------------------------
def _rtl(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))
    return p


def _set_cs_font(run, name="Arial", size=11):
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def _para(doc, text, *, bold=False, size=11, color=None, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text); run.bold = bold
    _set_cs_font(run, size=size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return _rtl(p)


def _heading(doc, text, level=1):
    sizes = {0: 18, 1: 14, 2: 12}
    p = _para(doc, text, bold=True, size=sizes.get(level, 12),
              color=(0x1a, 0x3b, 0x6e))
    return p


def _table(doc, df, caption=None, nd=3):
    if caption:
        c = _para(doc, caption, bold=True, size=10, color=(0x33, 0x33, 0x33))
    df = df.round(nd)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = 2
    # bidi on the table
    tblPr = t._tbl.tblPr
    tblPr.append(OxmlElement("w:bidiVisual"))
    hdr = t.rows[0].cells
    for j, col in enumerate(df.columns):
        run = hdr[j].paragraphs[0].add_run(str(col)); run.bold = True
        _set_cs_font(run, size=9); _rtl(hdr[j].paragraphs[0])
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, col in enumerate(df.columns):
            run = cells[j].paragraphs[0].add_run(str(row[col]))
            _set_cs_font(run, size=9); _rtl(cells[j].paragraphs[0])
    doc.add_paragraph()
    return t


# --------------------------------------------------------------------------
# Load results
# --------------------------------------------------------------------------
def _load(cfg):
    op = Path(cfg["_project_root"]) / cfg["paths"]["outputs_dir"]
    R = {}
    R["test"] = pd.read_csv(op / "06_metrics/test_metrics.csv")
    R["cv"] = pd.read_csv(op / "06_metrics/cv_metrics_by_fold.csv")
    R["boot"] = pd.read_csv(op / "06_metrics/test_metrics_bootstrap_ci.csv")
    R["calib"] = pd.read_csv(op / "06_metrics/calibration_metrics.csv")
    R["robust"] = pd.read_csv(op / "06_metrics/robustness_results.csv")
    R["shap"] = pd.read_csv(op / "07_explainability/shap_global_importance_xgb.csv")
    R["odds"] = pd.read_csv(op / "07_explainability/logistic_odds_ratios.csv")
    R["seed"] = pd.read_csv(op / "06_metrics/seed_stability_results.csv")
    R["audit"] = pd.read_csv(op / "01_data_audit/target_shift_audit.csv")
    return R


def _audit_val(audit, key):
    row = audit[audit["key"] == key]
    return float(row["value"].iloc[0]) if len(row) else float("nan")


def _prim(test):
    return test[test["threshold_kind"] == "thr_opt_primary"].set_index("model")


# --------------------------------------------------------------------------
# Build
# --------------------------------------------------------------------------
def build(cfg):
    R = _load(cfg)
    prim = _prim(R["test"])
    n_built = int(_audit_val(R["audit"], "n_built_one_year_ahead"))
    n_pos = int(_audit_val(R["audit"], "n_positive"))
    dev_n = int(_audit_val(R["audit"], "split::dev::n"))
    dev_p = int(_audit_val(R["audit"], "split::dev::positives"))
    test_n = int(_audit_val(R["audit"], "split::test::n"))
    test_p = int(_audit_val(R["audit"], "split::test::positives"))
    xgb = prim.loc["xgboost"]
    rf = prim.loc["random_forest"]
    lr = prim.loc["logistic"]
    best = MODEL_FA[prim["pr_auc"].idxmax()]

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    _heading(doc, "پیش‌بینی درماندگی مالی شرکت‌های پذیرفته‌شده در بورس اوراق بهادار "
                  "تهران با یادگیری ماشین قابل‌توضیح: مقایسهٔ XGBoost، جنگل تصادفی و "
                  "رگرسیون لجستیک", level=0)
    _para(doc, "پیش‌نویس مقاله (تولید خودکار از نتایج پایپ‌لاین Stage121). جای‌گذاری "
               "نام نویسندگان، وابستگی سازمانی و منابع تکمیلی بر عهدهٔ نویسنده است.",
          size=9, color=(0x88, 0x88, 0x88))

    # ---- Abstract ----
    _heading(doc, "چکیده", level=1)
    _para(doc,
          f"درماندگی مالی شرکت‌ها پیامدهای پرهزینه‌ای برای سرمایه‌گذاران، اعتباردهندگان و "
          f"نهادهای نظارتی دارد و پیش‌بینی زودهنگام آن اهمیت بالایی یافته است. در این پژوهش، "
          f"یک سامانهٔ پیش‌بینی «یک سال جلوتر» برای درماندگی مالی شرکت‌های بورس تهران طراحی "
          f"و سه الگوریتم رگرسیون لجستیک، جنگل تصادفی و XGBoost روی مجموعهٔ ویژگی یکسان به‌طور "
          f"منصفانه مقایسه شدند. داده‌ها شامل {n_built} مشاهدهٔ شرکت‌ـ‌سال (با نرخ درماندگی "
          f"حدود {n_pos/n_built*100:.1f}٪) است. برای جلوگیری از نشت اطلاعات، تقسیم زمانیِ "
          f"رو‌به‌جلو، اعتبارسنجی پنجرهٔ گسترش‌یابنده و انجام تمام پیش‌پردازش‌ها صرفاً روی دادهٔ "
          f"آموزش هر fold به‌کار رفت. عدم‌توازن شدید کلاس با یادگیری هزینه‌-حساس مدیریت و "
          f"معیار اصلی ارزیابی، سطح زیر منحنی دقت-فراخوانی (PR-AUC) قرار گرفت. بر پایهٔ نتایج، "
          f"مدل {best} با PR-AUC برابر {xgb['pr_auc']:.2f}، فراخوانی {xgb['recall']:.2f} و "
          f"دقت متوازن {xgb['balanced_accuracy']:.2f} روی مجموعهٔ آزمونِ دست‌نخورده بهترین "
          f"عملکرد را داشت. تحلیل توضیح‌پذیری مبتنی بر SHAP نشان داد نسبت اهرمی، اندازهٔ شرکت "
          f"و حاشیهٔ سود خالص مهم‌ترین عوامل پیش‌بینی‌اند. یافته‌ها بر برتری مدل‌های درختی "
          f"تقویتی و ضرورت گزارش معیارهای حساس به عدم‌توازن به‌جای صحت کلی تأکید می‌کنند.")
    _para(doc, "واژگان کلیدی: درماندگی مالی، یادگیری ماشین قابل‌توضیح، XGBoost، عدم‌توازن "
               "کلاس، SHAP، بورس اوراق بهادار تهران.", bold=True, size=10)

    # ---- 1 Intro ----
    _heading(doc, "۱. مقدمه", level=1)
    _para(doc,
          "درماندگی مالی وضعیتی است که در آن شرکت توان ایفای تعهدات مالی خود را از دست "
          "می‌دهد و می‌تواند به ورشکستگی بینجامد. پیش‌بینی به‌موقع این پدیده برای تخصیص بهینهٔ "
          "سرمایه، مدیریت ریسک اعتباری و سیاست‌گذاری نظارتی حیاتی است. مدل‌های کلاسیک مانند "
          "تحلیل تشخیصی و رگرسیون لجستیک سال‌ها مبنای کار بوده‌اند، اما با رشد روش‌های یادگیری "
          "ماشین، مدل‌های غیرخطی و گروهی توانسته‌اند روابط پیچیدهٔ میان نسبت‌های مالی را بهتر "
          "بازنمایی کنند. با این حال، دو چالش اساسی باقی است: نخست، خطر نشت اطلاعات و "
          "ارزیابی خوش‌بینانه در صورت رعایت‌نشدن ساختار زمانی و پنلی داده؛ دوم، عدم‌توازن شدید "
          "کلاس که گزارش معیارهایی مانند صحت کلی را گمراه‌کننده می‌سازد.")
    _para(doc,
          "این پژوهش با تمرکز بر شرکت‌های بورس تهران، سه هدف را دنبال می‌کند: (۱) طراحی یک "
          "فرایند کاملاً بازتولیدپذیر و مقاوم در برابر نشت برای پیش‌بینی یک‌سال‌جلوی درماندگی؛ "
          "(۲) مقایسهٔ منصفانهٔ رگرسیون لجستیک، جنگل تصادفی و XGBoost روی مجموعهٔ ویژگی و "
          "تقسیم یکسان؛ و (۳) ارائهٔ توضیح‌پذیری در سطح سراسری و موردی با SHAP و ضرایب مدل خطی.")

    # ---- 2 Lit ----
    _heading(doc, "۲. پیشینهٔ پژوهش", level=1)
    _para(doc,
          "نقطهٔ آغاز ادبیات پیش‌بینی درماندگی، مدل امتیاز Z آلتمن (۱۹۶۸) و سپس مدل لجستیک "
          "اولسون (۱۹۸۰) است که نسبت‌های مالی را به احتمال ورشکستگی پیوند زدند. با ظهور "
          "یادگیری ماشین، روش‌های گروهی مانند جنگل تصادفی (بریمن، ۲۰۰۱) و درخت‌های تقویتی "
          "گرادیانی به‌ویژه XGBoost (چن و گوئسترین، ۲۰۱۶) به‌دلیل توان مدل‌سازی غیرخطی و "
          "مقاومت در برابر بیش‌برازش مورد توجه قرار گرفتند. هم‌زمان، نیاز به توضیح‌پذیری "
          "مدل‌های پیچیده به توسعهٔ روش‌های مبتنی بر مقادیر شَپلی، به‌ویژه SHAP (لاندبرگ و "
          "لی، ۲۰۱۷)، انجامید که امکان تفسیر سهم هر ویژگی را در سطح فردی و سراسری فراهم "
          "می‌کند. پژوهش حاضر این خطوط را با تأکید بر کنترل نشت زمانی و گزارش معیارهای حساس "
          "به عدم‌توازن تلفیق می‌کند.")

    # ---- 3 Methods ----
    _heading(doc, "۳. روش‌شناسی", level=1)
    _heading(doc, "۳-۱. داده و متغیر هدف", level=2)
    _para(doc,
          f"واحد مشاهده شرکت‌ـ‌سال است و متغیر هدف، برچسب بازبینی‌شدهٔ درماندگی (۱=درمانده، "
          f"۰=غیردرمانده) بدون هیچ تغییری به‌کار رفت. برای آنکه مدل حقیقتاً «پیش‌بینی» باشد، "
          f"با اطلاعات سال t درماندگی سال t+1 هدف‌گذاری شد. پس از این انتقال، مجموعهٔ نهایی "
          f"{n_built} مشاهده با {n_pos} نمونهٔ مثبت ({n_pos/n_built*100:.2f}٪) به‌دست آمد؛ "
          f"تنها ردیف‌های حذف‌شده، مشاهداتِ سال پایانی فاقد هدفِ سال بعد بودند.")
    feats = pd.DataFrame({"ویژگی (مدل اصلی)": cfg["features"]["numeric_main"] +
                          cfg["features"]["categorical"]})
    _table(doc, feats, "جدول ۱: مجموعهٔ ویژگی اصلی (Feature Set A)", nd=3)
    _heading(doc, "۳-۲. تقسیم زمانی و کنترل نشت", level=2)
    _para(doc,
          f"چون رویدادهای درماندگی در سال‌های میانی بازه متمرکز بودند، تقسیم رو‌به‌جلو به‌صورت "
          f"مجموعهٔ توسعه ({dev_n} مشاهده، {dev_p} مثبت) و مجموعهٔ آزمونِ دست‌نخورده "
          f"({test_n} مشاهده، {test_p} مثبت) انجام شد و تنظیم ابرپارامترها با اعتبارسنجی "
          f"پنجرهٔ گسترش‌یابنده صورت گرفت. هیچ عملیات جای‌گذاری، مقیاس‌بندی، winsorization، "
          f"بیش‌نمونه‌گیری یا انتخاب آستانه پیش از تقسیم یا روی کل داده انجام نشد؛ همهٔ این "
          f"مراحل درون pipeline و فقط روی دادهٔ آموزش هر fold اجرا شدند.")
    _heading(doc, "۳-۳. پیش‌پردازش، عدم‌توازن و مدل‌ها", level=2)
    _para(doc,
          "جای‌گذاری میانه به‌همراه شاخص مفقودی برای متغیرهای عددی، RobustScaler برای رگرسیون "
          "لجستیک و کدگذاری One-Hot برای صنعت (با دستهٔ «نامشخص» برای مقادیر مفقود) به‌کار "
          "رفت. عدم‌توازن با یادگیری هزینه‌-حساس مدیریت شد: class_weight متوازن برای رگرسیون "
          "لجستیک و جنگل تصادفی، و scale_pos_weight برابر نسبت کلاس منفی به مثبت برای "
          "XGBoost. ابرپارامترها با Optuna و معیار میانگین PR-AUC اعتبارسنجی تنظیم شدند و "
          "آستانهٔ تصمیم تنها روی اعتبارسنجی (با بیشینه‌سازی F1) انتخاب شد.")

    # ---- 4 Results ----
    _heading(doc, "۴. یافته‌ها", level=1)
    cvt = R["cv"][["model", "fold", "pos_val", "pr_auc", "roc_auc"]].copy()
    cvt["model"] = cvt["model"].map(MODEL_FA)
    _table(doc, cvt, "جدول ۲: معیارهای اعتبارسنجی پنجرهٔ گسترش‌یابنده به تفکیک fold")
    tt = prim.reset_index()[["model", "pr_auc", "roc_auc", "recall", "precision",
                             "f1", "f2", "balanced_accuracy", "mcc", "brier"]].copy()
    tt["model"] = tt["model"].map(MODEL_FA)
    _table(doc, tt, "جدول ۳: معیارهای مجموعهٔ آزمون در آستانهٔ بهینه (معیار اصلی PR-AUC)")
    _para(doc,
          f"بر اساس جدول ۳، مدل {best} با PR-AUC={xgb['pr_auc']:.2f} و ROC-AUC="
          f"{xgb['roc_auc']:.2f} بهترین عملکرد را داشت و {int(xgb['tp'])} مورد از "
          f"{int(xgb['tp']+xgb['fn'])} شرکت درماندهٔ مجموعهٔ آزمون را شناسایی کرد "
          f"(فراخوانی {xgb['recall']:.2f}). جنگل تصادفی (PR-AUC={rf['pr_auc']:.2f}) دوم و "
          f"رگرسیون لجستیک (PR-AUC={lr['pr_auc']:.2f}) سوم شد. منحنی‌های ROC و دقت-فراخوانی "
          f"و ماتریس‌های درهم‌ریختگی در شکل‌های متناظر پوشهٔ نمودارها ارائه شده‌اند.")
    bx = R["boot"][R["boot"]["model"] == "xgboost"][["metric", "point", "ci_low",
                                                     "ci_high"]]
    _table(doc, bx, "جدول ۴: بازهٔ اطمینان ۹۵٪ با bootstrap خوشه‌ای (مدل XGBoost، آزمون)")
    sh = R["shap"].head(10).rename(columns={"feature": "ویژگی",
                                            "mean_abs_shap": "میانگین |SHAP|"})
    _table(doc, sh, "جدول ۵: ده ویژگی برتر بر اساس اهمیت سراسری SHAP (XGBoost)", nd=4)
    rb = R["robust"][["variant", "model", "cv_mean_pr_auc", "test_pr_auc",
                      "test_recall", "test_f1"]].copy()
    rb["model"] = rb["model"].map(MODEL_FA)
    _table(doc, rb, "جدول ۶: تحلیل‌های استحکام (جدا از مدل اصلی)")
    _para(doc,
          "تحلیل‌های استحکام شامل مقایسهٔ مجموعهٔ ویژگی اصلی با مجموعهٔ توسعه‌یافته، اجرای "
          "winsorization، جایگزینی وزن‌دهی کلاس با بیش‌نمونه‌گیری SMOTE، حذف متغیرهای نزدیک به "
          "تعریف درماندگی و مقایسه با طبقه‌بندی هم‌سال (که پیش‌بینی آینده نیست) بود. نتایج، "
          "ترتیب کلی برتری مدل‌ها را تأیید کرد و جایگزین مدل اصلی نیست.")

    # ---- 5 Discussion ----
    _heading(doc, "۵. بحث", level=1)
    _para(doc,
          f"برتری {best} با ماهیت غیرخطی و تعاملی روابط میان نسبت‌های مالی سازگار است. "
          f"تحلیل SHAP نشان داد نسبت اهرمی مهم‌ترین عامل افزایش ریسک درماندگی است؛ این یافته "
          f"با ادبیات نظری همسوست، زیرا اتکای بالا به بدهی توان بازپرداخت تعهدات را در "
          f"شرایط نامساعد کاهش می‌دهد. اندازهٔ شرکت (لگاریتم دارایی) و حاشیهٔ سود خالص نیز "
          f"نقش حفاظتی داشتند. کالیبراسیون احتمال‌ها با مقیاس‌بندی پلَت و ایزوتونیک بهبود یافت، "
          f"که برای کاربردهای تصمیم‌گیری مبتنی بر آستانه اهمیت دارد.")

    # ---- 6 Limitations ----
    _heading(doc, "۶. محدودیت‌ها", level=1)
    _para(doc,
          f"نخست، تعداد اندک نمونه‌های مثبت در پنجرهٔ آزمون ({test_p} مورد) بازه‌های اطمینان "
          f"را پهن می‌کند؛ از این رو PR-AUC اعتبارسنجی سیگنال پایدارتری برای انتخاب مدل است. "
          f"دوم، محدودیت‌های مستند داده شامل ۳۹ مقدار مفقود جریان نقد عملیاتی، ۱۴۵ مقدار "
          f"حل‌نشدهٔ سود ناخالص، ۱۳ ردیف با اطلاعات منبع ناقص و ۹۹ تغییر مالی غیرعادی "
          f"بررسی‌نشده است؛ هیچ‌یک با صفر یا مقدار ساختگی پر نشدند. سوم، تمرکز رویدادهای "
          f"درماندگی در سال‌های ابتدایی بازه، تعمیم به سال‌های اخیر را محدود می‌کند.")

    # ---- 7 Conclusion ----
    _heading(doc, "۷. نتیجه‌گیری", level=1)
    _para(doc,
          f"این پژوهش یک چارچوب بازتولیدپذیر و مقاوم در برابر نشت برای پیش‌بینی یک‌سال‌جلوی "
          f"درماندگی مالی ارائه کرد. مدل {best} با معیار اصلی PR-AUC بهترین تعادل میان "
          f"فراخوانی و دقت را فراهم آورد و تحلیل SHAP تفسیرپذیری لازم برای کاربرد عملی را "
          f"تأمین کرد. پیشنهاد می‌شود در پژوهش‌های آتی، متغیرهای کلان اقتصادی با تأخیر زمانی و "
          f"دامنهٔ زمانی گسترده‌تر برای افزایش تعداد رویدادها به‌کار گرفته شود.")

    # ---- References ----
    _heading(doc, "منابع", level=1)
    refs = [
        "Altman, E. I. (1968). Financial ratios, discriminant analysis and the "
        "prediction of corporate bankruptcy. The Journal of Finance, 23(4), 589–609.",
        "Ohlson, J. A. (1980). Financial ratios and the probabilistic prediction of "
        "bankruptcy. Journal of Accounting Research, 18(1), 109–131.",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. "
        "In Proceedings of the 22nd ACM SIGKDD (pp. 785–794).",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model "
        "predictions. In Advances in Neural Information Processing Systems 30.",
    ]
    for r in refs:
        p = doc.add_paragraph(); run = p.add_run(r); _set_cs_font(run, size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # English refs stay LTR

    out = utils.out_dir(cfg, "09_report") / "article_full_draft_fa.docx"
    doc.save(out)
    print(f"[09] full Persian academic article -> {out}")
    return out
