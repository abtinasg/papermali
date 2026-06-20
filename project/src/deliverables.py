"""Final publication deliverables: report PDF/DOCX, Persian article DOCX (RTL),
numbered article tables, figure index, analysis notebook, environment.yml.

Reads the produced CSVs under outputs/ so it can run standalone or from run_all.
"""
from __future__ import annotations
import json
from pathlib import Path

import pandas as pd

from . import utils


# ---------------------------------------------------------------------------
# Numbered, article-ready tables (Persian captions)
# ---------------------------------------------------------------------------
TABLE_SPECS = [
    ("Table_1", "جدول ۱: معیارهای مجموعهٔ آزمون سه مدل در آستانهٔ بهینهٔ مبتنی بر اعتبارسنجی",
     "06_metrics/test_metrics.csv",
     lambda df: df[df["threshold_kind"] == "thr_opt_primary"][
         ["model", "pr_auc", "roc_auc", "recall", "precision", "f1", "f2",
          "specificity", "balanced_accuracy", "mcc", "brier", "log_loss",
          "tp", "fp", "tn", "fn"]]),
    ("Table_2", "جدول ۲: معیارهای اعتبارسنجی پنجرهٔ گسترش‌یابنده به تفکیک fold",
     "06_metrics/cv_metrics_by_fold.csv",
     lambda df: df[["model", "fold", "n_val", "pos_val", "pr_auc", "roc_auc", "brier"]]),
    ("Table_3", "جدول ۳: بازهٔ اطمینان ۹۵٪ با bootstrap خوشه‌ای بر مبنای شرکت (آزمون)",
     "06_metrics/test_metrics_bootstrap_ci.csv",
     lambda df: df[["model", "metric", "point", "ci_low", "ci_high"]]),
    ("Table_4", "جدول ۴: پایداری معیارهای آزمون در ۳۰ seed (RF و XGBoost)",
     "06_metrics/seed_stability_summary.csv", lambda df: df),
    ("Table_5", "جدول ۵: تحلیل‌های استحکام (جدا از مدل اصلی)",
     "06_metrics/robustness_results.csv",
     lambda df: df[["variant", "model", "cv_mean_pr_auc", "test_pr_auc",
                    "test_recall", "test_precision", "test_f1", "test_brier"]]),
    ("Table_6", "جدول ۶: کالیبراسیون احتمال (Brier روی آزمون)",
     "06_metrics/calibration_metrics.csv", lambda df: df),
    ("Table_7", "جدول ۷: اهمیت سراسری SHAP — مدل XGBoost (آزمون)",
     "07_explainability/shap_global_importance_xgb.csv", lambda df: df.head(15)),
    ("Table_8", "جدول ۸: نسبت بخت و جهت اثر — رگرسیون لجستیک",
     "07_explainability/logistic_odds_ratios.csv", lambda df: df.head(20)),
]


def numbered_tables(cfg):
    out = utils.out_dir(cfg, "09_report")
    tdir = out / "tables"; tdir.mkdir(exist_ok=True)
    op = Path(cfg["_project_root"]) / cfg["paths"]["outputs_dir"]
    index = []
    with pd.ExcelWriter(out / "article_tables.xlsx") as xw:
        for tid, caption, rel, fn in TABLE_SPECS:
            src = op / rel
            if not src.exists():
                continue
            df = fn(pd.read_csv(src)).round(4)
            df.to_csv(tdir / f"{tid}.csv", index=False, encoding="utf-8-sig")
            cap = pd.DataFrame([[caption]], columns=[""])
            cap.to_excel(xw, sheet_name=tid, index=False, header=False, startrow=0)
            df.to_excel(xw, sheet_name=tid, index=False, startrow=2)
            index.append({"table": tid, "caption_fa": caption, "source": rel})
    pd.DataFrame(index).to_csv(out / "tables_index_fa.csv", index=False,
                               encoding="utf-8-sig")
    print(f"[09] numbered article tables -> {out/'article_tables.xlsx'}")


FIG_CAPTIONS = {
    "fig_class_distribution": "توزیع کلاس و نمونه‌های مثبت به تفکیک split و سال هدف",
    "fig_missing_data": "درصد داده‌های مفقود ویژگی‌ها به تفکیک split",
    "fig_roc_curves": "منحنی ROC سه مدل روی مجموعهٔ آزمون",
    "fig_precision_recall_curves": "منحنی دقت-فراخوانی (PR) سه مدل روی آزمون",
    "fig_confusion_matrices": "ماتریس درهم‌ریختگی در آستانهٔ بهینه (آزمون)",
    "fig_calibration_curves": "منحنی کالیبراسیون سه مدل روی آزمون",
    "fig_prob_distribution": "توزیع احتمال پیش‌بینی‌شده برای دو کلاس",
    "fig_threshold_performance": "عملکرد معیارها برحسب آستانه (روی اعتبارسنجی)",
    "fig_coefficient_plot": "ضرایب مهم رگرسیون لجستیک",
    "fig_feature_importance": "اهمیت ویژگی جنگل تصادفی و XGBoost",
    "fig_shap_bar_rf": "نمودار میله‌ای SHAP — جنگل تصادفی",
    "fig_shap_bar_xgb": "نمودار میله‌ای SHAP — XGBoost",
    "fig_shap_beeswarm_rf": "نمودار beeswarm SHAP — جنگل تصادفی",
    "fig_shap_beeswarm_xgb": "نمودار beeswarm SHAP — XGBoost",
}


def figure_index(cfg):
    out = utils.out_dir(cfg, "09_report")
    figs = utils.out_dir(cfg, "08_figures")
    stems = sorted({p.stem for p in figs.glob("*.png")})
    recs, n = [], 0
    # keep a stable, meaningful order: main figs first, then shap detail
    order = list(FIG_CAPTIONS.keys())
    extra = [s for s in stems if s not in FIG_CAPTIONS]
    for s in order + extra:
        if s not in stems:
            continue
        n += 1
        cap = FIG_CAPTIONS.get(s)
        if cap is None:
            if "dependence" in s:
                cap = "نمودار وابستگی SHAP: " + s.split("dependence_")[-1]
            elif "waterfall" in s:
                cap = "نمودار waterfall SHAP برای یک نمونه: " + s.split("waterfall_")[-1]
            else:
                cap = s
        recs.append({"figure_no": n, "stem": s, "png": f"08_figures/{s}.png",
                     "pdf": f"08_figures/{s}.pdf", "caption_fa": cap})
    pd.DataFrame(recs).to_csv(out / "figure_index_fa.csv", index=False,
                              encoding="utf-8-sig")
    print(f"[09] figure index ({n} figures) -> {out/'figure_index_fa.csv'}")


# ---------------------------------------------------------------------------
# Technical report -> PDF (English content renders cleanly)
# ---------------------------------------------------------------------------
def report_pdf(cfg):
    import markdown
    from xhtml2pdf import pisa
    out = utils.out_dir(cfg, "09_report")
    md_path = out / "technical_modeling_report.md"
    if not md_path.exists():
        return
    html_body = markdown.markdown(md_path.read_text(encoding="utf-8"),
                                  extensions=["tables", "fenced_code"])
    css = """<style>
      body{font-family:Helvetica,Arial,sans-serif;font-size:9px;color:#111}
      h1{font-size:16px} h2{font-size:12px;color:#1a3b6e;margin-top:12px}
      table{border-collapse:collapse;width:100%;margin:6px 0}
      th,td{border:1px solid #999;padding:3px;font-size:8px}
      th{background:#e8eef7}
    </style>"""
    html = f"<html><head><meta charset='utf-8'>{css}</head><body>{html_body}</body></html>"
    with open(out / "technical_modeling_report.pdf", "wb") as f:
        pisa.CreatePDF(html, dest=f)
    print(f"[09] technical report PDF -> {out/'technical_modeling_report.pdf'}")


# ---------------------------------------------------------------------------
# Word documents (python-docx). Persian paragraphs set to RTL.
# ---------------------------------------------------------------------------
def _set_rtl(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _md_to_docx(doc, md_text, rtl=False):
    """Minimal markdown renderer: #/##/### headings, bullets, paragraphs.
    Markdown tables are skipped here (added separately as real Word tables)."""
    in_table = False
    for line in md_text.splitlines():
        s = line.rstrip()
        if s.startswith("|"):  # skip raw md tables
            in_table = True
            continue
        if in_table and not s.startswith("|"):
            in_table = False
        if not s:
            continue
        if s.startswith("### "):
            p = doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            p = doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            p = doc.add_heading(s[2:], level=1)
        elif s.startswith("- "):
            p = doc.add_paragraph(s[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(s)
        if rtl:
            _set_rtl(p)


def _add_df_table(doc, df, rtl=False):
    df = df.round(4)
    t = doc.add_table(rows=1, cols=len(df.columns)); t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].text = str(c)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(row[c])
    if rtl:
        t.alignment = 2


def report_docx(cfg):
    from docx import Document
    out = utils.out_dir(cfg, "09_report")
    md_path = out / "technical_modeling_report.md"
    if not md_path.exists():
        return
    doc = Document()
    _md_to_docx(doc, md_path.read_text(encoding="utf-8"), rtl=False)
    # append the key test-metrics table as a real Word table
    op = Path(cfg["_project_root"]) / cfg["paths"]["outputs_dir"]
    tm = pd.read_csv(op / "06_metrics/test_metrics.csv")
    tm = tm[tm["threshold_kind"] == "thr_opt_primary"][
        ["model", "pr_auc", "roc_auc", "recall", "precision", "f1",
         "balanced_accuracy", "mcc", "brier"]]
    doc.add_heading("Appendix table — test metrics (primary threshold)", level=2)
    _add_df_table(doc, tm)
    doc.save(out / "technical_modeling_report.docx")
    print(f"[09] technical report DOCX -> {out/'technical_modeling_report.docx'}")


def article_docx(cfg):
    from docx import Document
    out = utils.out_dir(cfg, "09_report")
    op = Path(cfg["_project_root"]) / cfg["paths"]["outputs_dir"]
    # methods
    mp = out / "article_methods_draft_fa.md"
    if mp.exists():
        doc = Document()
        _md_to_docx(doc, mp.read_text(encoding="utf-8"), rtl=True)
        doc.save(out / "article_methods_draft_fa.docx")
    # results (+ real metrics table)
    rp = out / "article_results_draft_fa.md"
    if rp.exists():
        doc = Document()
        _md_to_docx(doc, rp.read_text(encoding="utf-8"), rtl=True)
        tm = pd.read_csv(op / "06_metrics/test_metrics.csv")
        tm = tm[tm["threshold_kind"] == "thr_opt_primary"][
            ["model", "pr_auc", "roc_auc", "recall", "precision", "f1",
             "balanced_accuracy", "brier"]]
        h = doc.add_heading("جدول ۱: معیارهای آزمون", level=2); _set_rtl(h)
        _add_df_table(doc, tm, rtl=True)
        doc.save(out / "article_results_draft_fa.docx")
    print(f"[09] Persian article DOCX (methods + results) saved")


# ---------------------------------------------------------------------------
# Analysis notebook + environment.yml
# ---------------------------------------------------------------------------
def analysis_notebook(cfg):
    out = utils.out_dir(cfg, "03_code")
    cells = [
        ("md", "# Analysis notebook — Financial Distress (Stage121)\n"
               "Display-only. Run `python run_all.py` first to generate `outputs/`."),
        ("code", "import pandas as pd, json\nfrom pathlib import Path\n"
                 "OUT = Path('..')  # this notebook lives in outputs/03_code; '..' = outputs/"),
        ("md", "## Test metrics (primary threshold)"),
        ("code", "tm = pd.read_csv(OUT/'06_metrics/test_metrics.csv')\n"
                 "tm[tm.threshold_kind=='thr_opt_primary']"),
        ("md", "## Cross-validation by fold"),
        ("code", "pd.read_csv(OUT/'06_metrics/cv_metrics_by_fold.csv')"),
        ("md", "## Bootstrap 95% CIs"),
        ("code", "pd.read_csv(OUT/'06_metrics/test_metrics_bootstrap_ci.csv')"),
        ("md", "## Robustness"),
        ("code", "pd.read_csv(OUT/'06_metrics/robustness_results.csv')"),
        ("md", "## SHAP global importance (XGBoost) + figures"),
        ("code", "pd.read_csv(OUT/'07_explainability/shap_global_importance_xgb.csv').head(15)"),
        ("code", "from IPython.display import Image\n"
                 "Image(str(OUT/'08_figures/fig_precision_recall_curves.png'))"),
    ]
    nb = {"cells": [], "metadata": {"kernelspec": {"name": "python3",
          "display_name": "Python 3"}}, "nbformat": 4, "nbformat_minor": 5}
    for kind, src in cells:
        if kind == "md":
            nb["cells"].append({"cell_type": "markdown", "metadata": {},
                                "source": src.splitlines(keepends=True)})
        else:
            nb["cells"].append({"cell_type": "code", "metadata": {}, "outputs": [],
                                "execution_count": None,
                                "source": src.splitlines(keepends=True)})
    (out / "analysis_display.ipynb").write_text(json.dumps(nb, ensure_ascii=False,
                                                           indent=1), encoding="utf-8")
    print(f"[03] analysis notebook -> {out/'analysis_display.ipynb'}")


def environment_yml(cfg):
    repro = utils.out_dir(cfg, "10_reproducibility")
    content = (
        "name: financial-distress-stage121\n"
        "channels: [conda-forge]\n"
        "dependencies:\n"
        "  - python=3.13\n"
        "  - pip\n"
        "  - pip:\n"
        "      - -r ../../requirements.txt\n")
    (repro / "environment.yml").write_text(content, encoding="utf-8")
    # also at project root for convenience
    (Path(cfg["_project_root"]) / "environment.yml").write_text(content, encoding="utf-8")


def copy_code(cfg):
    """Make outputs/03_code a self-contained copy of the runnable code (brief 18)."""
    import shutil
    out = utils.out_dir(cfg, "03_code")
    root = Path(cfg["_project_root"])
    if (out / "src").exists():
        shutil.rmtree(out / "src")
    shutil.copytree(root / "src", out / "src",
                    ignore=shutil.ignore_patterns("__pycache__"))
    for f in ["run_all.py", "config.yaml", "requirements.txt", "README_RUN.md",
              "environment.yml"]:
        if (root / f).exists():
            shutil.copy(root / f, out / f)
    print(f"[03] runnable code copied -> {out}")


def run(cfg):
    numbered_tables(cfg)
    figure_index(cfg)
    report_pdf(cfg)
    report_docx(cfg)
    article_docx(cfg)
    from . import article, outputs_report
    article.build(cfg)
    outputs_report.build(cfg)
    analysis_notebook(cfg)
    environment_yml(cfg)
    copy_code(cfg)
